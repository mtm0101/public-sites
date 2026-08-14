from __future__ import annotations

import json
import html
import re
import shutil
import sys
import unicodedata
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from dol_vocab_utils import enrich_doc_examples

try:
    from deep_translator import GoogleTranslator
except Exception:  # pragma: no cover - optional local publishing helper
    GoogleTranslator = None


SOURCE_FILES = {
    "vocab_final_words": "02c Thien Speaking IELTS Fighter 2024 Final Words.docx",
    "vocab_examples": "02a2 IELTS_Speaking_Q2_Vocabulary_Examples_Read_Aloud.docx",
    "speaking_ai": "02b Thien Speaking IELTS Fighter 2024 - Simplification by AI - 20250204.docx",
    "speaking_groups": "02a1 IELTS_Speaking_Q2_English_Groups_Questions_Bilingual_Samples.docx",
}


def clean(text: str) -> str:
    return " ".join((text or "").replace("\u00a0", " ").split())


def slug(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = "".join(ch for ch in normalized if not unicodedata.combining(ch)).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "-", ascii_text.lower()).strip("-")[:90]


def iter_blocks(parent: DocumentType | _Cell):
    parent_element = parent.element.body if isinstance(parent, DocumentType) else parent._tc
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def image_targets(paragraph: Paragraph) -> list[str]:
    rel_ids = paragraph._p.xpath(".//a:blip/@r:embed")
    targets = []
    for rel_id in rel_ids:
        rel = paragraph.part.rels.get(rel_id)
        if rel is not None:
            targets.append(str(rel.target_ref).replace("\\", "/"))
    return targets


def anonymize_sample(text: str) -> str:
    replacements = [
        (r"my friend, Hung, a pediatrician", "a close friend who is a pediatrician"),
        (r"my friend, Hung, who", "a close friend who"),
        (r"my friend, Linh, who", "a close friend who"),
        (r"my friend, Linh, a university lecturer", "a close friend who is a university lecturer"),
        (r"my friend, Nhat, who", "a close friend who"),
        (r"my old friend, Nhat, whom", "an old friend whom"),
        (r"Elon Musk, a renowned entrepreneur and innovator", "a renowned entrepreneur and innovator"),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.I)
    return text


def copy_docx_asset(docx_path: Path, target_ref: str, assets_dir: Path, prefix: str) -> str:
    member = "word/" + target_ref.lstrip("/")
    filename = f"{prefix}-{Path(target_ref).name}"
    destination = assets_dir / filename
    with zipfile.ZipFile(docx_path) as archive:
        with archive.open(member) as source, destination.open("wb") as output:
            shutil.copyfileobj(source, output)
    return f"data/codex/speaking/assets/{assets_dir.name}/{filename}"


def extract_vocab(docx_path: Path, output_dir: Path, import_time: str) -> Path:
    doc = Document(docx_path)
    section = "Vocabulary"
    section_num = 0
    items = []
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            style = block.style.name if block.style else ""
            if text and style == "Section Heading":
                section = text
                section_num += 1
            continue
        if not block.rows or not block.rows[0].cells:
            continue
        paragraphs = [clean(p.text) for p in block.rows[0].cells[0].paragraphs if clean(p.text)]
        if len(paragraphs) < 2:
            continue
        headword, example = paragraphs[0], paragraphs[-1]
        items.append({
            "id": f"v-{len(items)+1:03d}-{slug(headword)}",
            "passageLabel": section,
            "passageName": section,
            "passageNum": section_num,
            "text": headword,
            "pos": "",
            "vn": "",
            "ipaUK": "",
            "example": example,
            "exampleVi": "",
            "exampleIpa": "",
            "l1": "Others",
            "l2": "Speaking 2026",
            "l3": docx_path.stem,
            "l4": section,
        })

    items = clean_vocab_items(items)

    # Persist translations in the generated JSON. Sentence IPA is supplied by
    # index.html at render time, but Vietnamese should survive cache clears.
    enrich_doc_examples({"items": items}, translate=True, rebuild_html=False, progress=True)
    for item in items:
        if not (item.get("vn") or "").strip():
            item["vn"] = translate_speaking_vi(item.get("text") or "")
            item["meaningVi"] = item["vn"]

    identifier = f"dol-gpt-codex-speaking-2026-{slug(docx_path.stem)}"
    clean_title = re.sub(r"[_\\]+", " ", docx_path.stem).strip()
    by_section: dict[str, list[dict]] = {}
    for item in items:
        by_section.setdefault(item["passageLabel"], []).append(item)
    detail_sections = []
    for index, (label, section_items) in enumerate(by_section.items(), 1):
        cards = []
        for item in section_items:
            cards.append(
                '<div class="vocab dol-card">'
                f'<div class="head">{html.escape(item["text"])} <span class="ipa"></span></div>'
                f'<p class="vi main-meaning">{html.escape(item.get("vn") or "")}</p>'
                f'<p class="en ex">{html.escape(item["example"])}</p>'
                f'<p class="vi ex-vi">{html.escape(item["exampleVi"])}</p>'
                '</div>'
            )
        detail_sections.append({
            "id": f"vocab-{index:02d}-{slug(label)}",
            "title": label.title(),
            "level": 1,
            "html": "\n".join(cards),
        })
    detail_sections.append({
        "id": "source",
        "title": "Source",
        "level": 1,
        "html": (
            "<p class=\"en\">Vocabulary and examples extracted from the source Word document for personal IELTS speaking study.</p>"
            "<p class=\"vn\">Từ vựng và ví dụ được trích xuất từ tài liệu Word nguồn để phục vụ việc học IELTS Speaking cá nhân.</p>"
        ),
    })

    payload = {
        "schema": 2,
        "format": "dol-vocab",
        "id": identifier,
        "type": "vocab",
        "source": "codex",
        "category": "dol-vocab",
        "topicNumber": 0,
        "title": clean_title,
        "group": "Others",
        "book": "Speaking 2026",
        "bookNum": 2026,
        "testNum": 0,
        "skill": "speaking",
        "superlmsFlat": True,
        "l1": "Others",
        "l2": "Speaking 2026",
        "l3": clean_title,
        "dateTime": "2026-07-17T11:00",
        "contentUpdatedAt": import_time,
        "wordCount": len(items),
        "words": [item["text"] for item in items],
        "items": items,
        "sections": detail_sections,
    }
    output = output_dir / f"{identifier}.json"
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=1), encoding="utf-8")
    return output


def norm_vocab_key(text: str) -> str:
    return re.sub(r"\s+", " ", clean(text).lower().replace("…", "...")).strip(" .!?")


def headword_is_promoted_example(text: str) -> bool:
    text = clean(text)
    if not text:
        return False
    if "…" in text or "..." in text:
        return False
    words = text.split()
    return bool(re.search(r"[.!?][\"'”’)]*$", text)) and (len(words) >= 5 or "," in text or ";" in text)


def clean_vocab_items(items: list[dict]) -> list[dict]:
    """Remove exact duplicate headwords and full-sentence examples promoted to headwords."""
    by_key: dict[str, dict] = {}
    cleaned: list[dict] = []
    promoted_examples: list[dict] = []

    for item in items:
        key = norm_vocab_key(item.get("text") or "")
        if not key:
            continue
        if key in by_key:
            continue
        if headword_is_promoted_example(item.get("text") or ""):
            promoted_examples.append(item)
            continue
        by_key[key] = item
        cleaned.append(item)

    # If a removed sentence clearly demonstrates an existing template headword,
    # keep it as that template's example instead of losing useful study content.
    for removed in promoted_examples:
        text = clean(removed.get("text") or "")
        text_key = norm_vocab_key(text)
        best = None
        best_len = 0
        for key, candidate in by_key.items():
            cand_text = clean(candidate.get("text") or "")
            if "…" not in cand_text and "..." not in cand_text:
                continue
            prefix = norm_vocab_key(cand_text).replace("...", "").strip(" ,")
            if prefix and prefix in text_key and len(prefix) > best_len:
                best = candidate
                best_len = len(prefix)
        if best:
            if best.get("_cleanPromotedExample"):
                continue
            best["example"] = removed.get("example") or text
            if removed.get("exampleVi"):
                best["exampleVi"] = removed["exampleVi"]
            best["_cleanPromotedExample"] = True

    for index, item in enumerate(cleaned, 1):
        item.pop("_cleanPromotedExample", None)
        item["id"] = f"v-{index:03d}-{slug(item.get('text') or '')}"
    return cleaned


@dataclass
class SpeakingSample:
    heading: str
    group: str = ""
    source_no: int | None = None
    questions: list[dict] = field(default_factory=list)
    segments: list[dict] = field(default_factory=list)


_SPEAKING_TRANSLATOR = None
_SPEAKING_TRANSLATE_CACHE: dict[str, str] = {}


def translate_speaking_vi(text: str) -> str:
    text = clean(text)
    if not text or GoogleTranslator is None:
        return ""
    key = text[:480]
    if key in _SPEAKING_TRANSLATE_CACHE:
        return _SPEAKING_TRANSLATE_CACHE[key]
    global _SPEAKING_TRANSLATOR
    if _SPEAKING_TRANSLATOR is None:
        _SPEAKING_TRANSLATOR = GoogleTranslator(source="en", target="vi")
    try:
        out = clean(_SPEAKING_TRANSLATOR.translate(text) or "")
    except Exception:
        out = ""
    if out and out.lower() != text.lower():
        _SPEAKING_TRANSLATE_CACHE[key] = out
        return out
    return ""


def question_record(text: str, number: int | None = None) -> dict:
    rec = {"text": clean(text)}
    if number:
        rec["number"] = number
    vi = translate_speaking_vi(rec["text"])
    if vi:
        rec["vi"] = vi
    return rec


def enrich_speaking_segments(sample: SpeakingSample) -> None:
    for segment in sample.segments:
        segment.pop("target", None)
        segment.pop("vietnamese_image", None)
        if not segment.get("vietnamese"):
            vi = translate_speaking_vi(segment.get("english") or "")
            if vi:
                segment["vietnamese"] = vi


def extract_ai_speaking(docx_path: Path, assets_dir: Path) -> list[SpeakingSample]:
    doc = Document(docx_path)
    samples: list[SpeakingSample] = []
    current: SpeakingSample | None = None
    current_group = ""
    in_answer = False

    for block in iter_blocks(doc):
        if isinstance(block, Table):
            if current and block.rows and block.rows[0].cells:
                text = clean(block.rows[0].cells[0].text)
                text = re.sub(r"^QUESTION\s*", "", text, flags=re.I)
                if text:
                    current.questions = [question_record(text, current.source_no)]
            continue

        text = clean(block.text)
        style = block.style.name if block.style else ""
        if style == "Heading 1":
            in_answer = False
            if text.upper() == "TABLE OF CONTENTS":
                current = None
            else:
                current_group = text
            continue
        if style == "Heading 2":
            m = re.match(r"^(\d+)\.\s*(.*)$", text)
            heading = (m.group(2) if m else re.sub(r"^\d+\.\s*", "", text)).title()
            current = SpeakingSample(heading=heading, group=current_group, source_no=int(m.group(1)) if m else None)
            samples.append(current)
            in_answer = False
            continue
        if not current:
            continue
        if text.upper() == "SAMPLE ANSWER":
            in_answer = True
            continue
        if not in_answer:
            continue
        targets = image_targets(block)
        if text:
            current.segments.append({"english": anonymize_sample(text), "target": ""})
        for target in targets:
            if current.segments:
                current.segments[-1]["target"] = target

    for sample in samples:
        enrich_speaking_segments(sample)
    return samples


def extract_group_speaking(docx_path: Path, assets_dir: Path) -> list[SpeakingSample]:
    doc = Document(docx_path)
    samples: list[SpeakingSample] = []
    current: SpeakingSample | None = None
    collecting_questions = False

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            style = block.style.name if block.style else ""
            if style == "Heading 1" and text.upper().startswith("GROUP "):
                heading = re.sub(r"^GROUP\s+\d+\s*:\s*", "", text, flags=re.I).title()
                current = SpeakingSample(heading=heading, group=heading)
                samples.append(current)
                collecting_questions = False
            elif current and style == "Heading 2" and text.startswith("Questions"):
                collecting_questions = True
            elif current and style == "Heading 2" and text.startswith("Sample Answer"):
                collecting_questions = False
            elif current and collecting_questions and style == "List Bullet" and text:
                m = re.match(r"^(\d+)\.\s*(.*)$", text)
                current.questions.append(question_record(m.group(2) if m else text, int(m.group(1)) if m else None))
            continue

        if not current or not block.rows or not block.rows[0].cells:
            continue
        segments = []
        for paragraph in block.rows[0].cells[0].paragraphs:
            text = clean(paragraph.text)
            if text:
                segments.append({"english": text, "target": ""})
            for target in image_targets(paragraph):
                if segments:
                    segments[-1]["target"] = target
        if segments:
            current.segments = segments

    for sample in samples:
        enrich_speaking_segments(sample)
    return samples


def write_speaking_source(
    docx_path: Path,
    samples: list[SpeakingSample],
    prefix: str,
    output_dir: Path,
    import_time: str,
) -> Path:
    items = []
    for item_no, sample in enumerate(samples, 1):
        topic = sample.heading
        questions = []
        question_entries = [(question_no, question, [question]) for question_no, question in enumerate(sample.questions, 1)]
        if prefix == "Ngoc Bach" and sample.questions:
            question_entries = [(1, sample.questions[0], sample.questions)]
        for question_no, question, related_questions in question_entries:
            question_text = question.get("text") if isinstance(question, dict) else clean(str(question))
            english = "\n\n".join(segment["english"] for segment in sample.segments)
            q_item = {
                "question_no": question_no,
                "question": question_text,
                "question_vi": question.get("vi", "") if isinstance(question, dict) else "",
                "answers": [{
                    "label": "Answer",
                    "english": english,
                    "vietnamese": "",
                    "segments": sample.segments,
                }],
            }
            if related_questions:
                q_item["related_questions"] = related_questions
            questions.append(q_item)
        items.append({
            "item_no": item_no,
            "topic": topic,
            "title": f"{topic} — Speaking Part 2 sample",
            "l1": prefix,
            "l2": sample.group or sample.heading,
            "questions": questions,
        })

    identifier = f"dol-speaking-codex-speaking-2026-{slug(docx_path.stem)}"
    payload = {
        "schema": 2,
        "format": "speaking-source",
        "id": identifier,
        "type": "speaking",
        "source": "codex",
        "category": "dol-speaking",
        "title": docx_path.stem,
        "l1": prefix,
        "l2": docx_path.stem,
        "dateTime": "2026-07-17T11:00",
        "contentUpdatedAt": import_time,
        "items": items,
    }
    output = output_dir / f"{identifier}.json"
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=1), encoding="utf-8")
    return output


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: import_speaking_2026_docx.py SOURCE_FOLDER TOOL_ROOT")
    source_folder = Path(sys.argv[1]).resolve()
    tool_root = Path(sys.argv[2]).resolve()
    vocab_output = tool_root / "data" / "codex" / "dol"
    speaking_output = tool_root / "data" / "codex" / "speaking"
    assets_dir = speaking_output / "assets" / "speaking-2026"
    vocab_output.mkdir(parents=True, exist_ok=True)
    speaking_output.mkdir(parents=True, exist_ok=True)
    assets_dir.mkdir(parents=True, exist_ok=True)
    import_time = datetime.now().astimezone().isoformat(timespec="seconds")

    outputs = [
        extract_vocab(source_folder / SOURCE_FILES["vocab_final_words"], vocab_output, import_time),
        extract_vocab(source_folder / SOURCE_FILES["vocab_examples"], vocab_output, import_time),
    ]
    ai_doc = source_folder / SOURCE_FILES["speaking_ai"]
    group_doc = source_folder / SOURCE_FILES["speaking_groups"]
    outputs.append(write_speaking_source(
        ai_doc, extract_ai_speaking(ai_doc, assets_dir), "My AI", speaking_output, import_time
    ))
    outputs.append(write_speaking_source(
        group_doc, extract_group_speaking(group_doc, assets_dir), "Ngoc Bach", speaking_output, import_time
    ))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
